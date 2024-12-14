from docx import Document
from docx.shared import Cm


def main():
    doc = Document("output.docx")

    doc.add_picture("picture.jpg", width=Cm(8), height=Cm(8))
    doc.add_paragraph("Подпись: человечек")
    
    doc.save("output.docx")


if __name__ == "__main__":
    main()
